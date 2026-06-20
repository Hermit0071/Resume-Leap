import json
import os
import shutil
import subprocess
import sys
import tempfile
from html import escape as html_escape
from datetime import datetime
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import quote, unquote, urlparse

RUNTIME_PYTHON_LIB = Path("/Users/lwc/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/lib")
if RUNTIME_PYTHON_LIB.exists():
    for candidate in sorted(RUNTIME_PYTHON_LIB.glob("python*/site-packages")):
        sys.path.insert(0, str(candidate))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
STATIC_DIR = ROOT / "web"
EXPORT_DIR = ROOT / "exports"
HOST = "127.0.0.1"
PORT = 8123
CHROME_RENDERER = ROOT / "chrome_exporter.mjs"
NODE_CANDIDATES = [
    Path("/Users/lwc/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node"),
    Path(shutil.which("node") or ""),
]

FONT_CANDIDATES = [
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/STHeiti Light.ttc",
]

DOCX_FONT = "PingFang SC"
REPORTLAB_FONT_NAME = "STSong-Light"


def ensure_export_dir() -> None:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def find_node_executable() -> Path | None:
    for candidate in NODE_CANDIDATES:
        if candidate and candidate.is_file() and os.access(candidate, os.X_OK):
            return candidate
    return None


def find_font_path() -> str | None:
    for candidate in FONT_CANDIDATES:
        if os.path.exists(candidate):
            return candidate
    return None


FONT_PATH = find_font_path()
try:
    pdfmetrics.registerFont(UnicodeCIDFont(REPORTLAB_FONT_NAME))
except Exception:
    REPORTLAB_FONT_NAME = "Helvetica"


def sanitize_filename(name: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in "-_ " else "_" for ch in name).strip()
    return cleaned or "resume"


def normalize_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def normalize_list(items) -> list[str]:
    if not items:
        return []
    return [normalize_text(item) for item in items if normalize_text(item)]


def normalize_section(section: dict) -> dict:
    section_type = section.get("type", "list")
    items = section.get("items", [])
    if section_type == "text":
        text_value = ""
        if items and isinstance(items[0], dict):
            text_value = items[0].get("text", "")
        else:
            text_value = section.get("text", "")
        normalized_items = [{"text": normalize_text(text_value)}]
    else:
        normalized_items = []
        for item in items:
            normalized_items.append(
                {
                    "title": normalize_text(item.get("title", "")),
                    "subtitle": normalize_text(item.get("subtitle", "")),
                    "meta": normalize_text(item.get("meta", "")),
                    "detail": normalize_text(item.get("detail", "")),
                    "bullets": normalize_list(item.get("bullets", [])),
                }
            )
    return {
        "id": normalize_text(section.get("id", "")),
        "title": normalize_text(section.get("title", "未命名板块")),
        "type": section_type,
        "visible": bool(section.get("visible", True)),
        "items": normalized_items,
    }


def build_safe_stem(name: str) -> str:
    base = sanitize_filename(name)
    candidate = EXPORT_DIR / base
    if not any(EXPORT_DIR.glob(f"{base}*")):
        return base
    return f"{base}-{datetime.now().strftime('%Y%m%d-%H%M%S-%f')}"


def pdf_text(value: str) -> str:
    return html_escape(value).replace("\n", "<br/>")


def pdf_inline(value: str) -> str:
    return html_escape(value)


def normalize_payload(payload: dict) -> dict:
    personal = payload.get("personal", {})
    theme = payload.get("theme", {})
    sections = [normalize_section(section) for section in payload.get("sections", []) if section.get("visible", True)]
    return {
        "templateId": normalize_text(payload.get("templateId", "harbor")),
        "theme": {
            "accent": normalize_text(theme.get("accent", "#1f4e79")) or "#1f4e79",
            "fontScale": max(0.9, min(1.2, float(theme.get("fontScale", 1.0) or 1.0))),
            "density": normalize_text(theme.get("density", "balanced")) or "balanced",
        },
        "personal": {
            "name": normalize_text(personal.get("name", "你的姓名")),
            "tagline": normalize_text(personal.get("tagline", "高中申请者")),
            "school": normalize_text(personal.get("school", "")),
            "email": normalize_text(personal.get("email", "")),
            "phone": normalize_text(personal.get("phone", "")),
            "location": normalize_text(personal.get("location", "")),
            "website": normalize_text(personal.get("website", "")),
        },
        "sections": sections,
    }


def set_run_font(run, font_name: str, size: float | None = None, color: str | None = None, bold: bool | None = None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    if bold is not None:
        run.bold = bold


def add_bordered_heading(document: Document, text: str, accent: str):
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), accent.replace("#", ""))
    border.append(bottom)
    p_pr.append(border)
    run = paragraph.add_run(text)
    set_run_font(run, DOCX_FONT, 13, accent, True)


def build_docx(data: dict, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    personal = data["personal"]
    accent = data["theme"]["accent"]
    scale = data["theme"]["fontScale"]

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(personal["name"])
    set_run_font(title_run, DOCX_FONT, 21 * scale, None, True)

    if personal["tagline"]:
        tagline = document.add_paragraph()
        tagline_run = tagline.add_run(personal["tagline"])
        set_run_font(tagline_run, DOCX_FONT, 11.5 * scale, accent)

    contacts = " | ".join(filter(None, [personal["school"], personal["email"], personal["phone"], personal["location"], personal["website"]]))
    if contacts:
        meta = document.add_paragraph()
        meta_run = meta.add_run(contacts)
        set_run_font(meta_run, DOCX_FONT, 10 * scale, "4A4A4A")

    for section_data in data["sections"]:
        add_bordered_heading(document, section_data["title"], accent)
        if section_data["type"] == "text":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(section_data["items"][0]["text"])
            set_run_font(run, DOCX_FONT, 10.5 * scale)
            paragraph.paragraph_format.space_after = Pt(6)
            continue

        for item in section_data["items"]:
            if item["title"]:
                line = document.add_paragraph()
                line.paragraph_format.space_after = Pt(0)
                title_run = line.add_run(item["title"])
                set_run_font(title_run, DOCX_FONT, 11 * scale, None, True)
                if item["subtitle"]:
                    subtitle_run = line.add_run(f"  {item['subtitle']}")
                    set_run_font(subtitle_run, DOCX_FONT, 10.5 * scale, "555555")
                if item["meta"]:
                    meta_run = line.add_run(f"   {item['meta']}")
                    set_run_font(meta_run, DOCX_FONT, 10 * scale, accent)
            if item["detail"]:
                detail = document.add_paragraph()
                detail.paragraph_format.left_indent = Cm(0.3)
                detail_run = detail.add_run(item["detail"])
                set_run_font(detail_run, DOCX_FONT, 10.2 * scale, "333333")
                detail.paragraph_format.space_after = Pt(0)
            for bullet in item["bullets"]:
                bullet_p = document.add_paragraph(style="List Bullet")
                bullet_p.paragraph_format.left_indent = Cm(0.7)
                bullet_run = bullet_p.add_run(bullet)
                set_run_font(bullet_run, DOCX_FONT, 10 * scale, "333333")

    document.save(str(output_path))


def build_pdf(data: dict, output_path: Path) -> None:
    personal = data["personal"]
    accent = data["theme"]["accent"]
    scale = data["theme"]["fontScale"]
    font_name = REPORTLAB_FONT_NAME
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=22 * scale,
        textColor=HexColor("#111111"),
        leading=26 * scale,
        spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "ResumeSubtitle",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11 * scale,
        textColor=HexColor(accent),
        leading=14 * scale,
        spaceAfter=3,
    )
    meta_style = ParagraphStyle(
        "ResumeMeta",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9.5 * scale,
        textColor=HexColor("#555555"),
        leading=12 * scale,
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=12 * scale,
        textColor=HexColor(accent),
        leading=14 * scale,
        borderPadding=0,
        spaceBefore=8,
        spaceAfter=5,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10 * scale,
        textColor=HexColor("#222222"),
        leading=14 * scale,
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=body_style,
        leftIndent=10,
        bulletIndent=0,
    )

    story.append(Paragraph(personal["name"], title_style))
    if personal["tagline"]:
        story.append(Paragraph(pdf_text(personal["tagline"]), subtitle_style))
    contacts = " | ".join(filter(None, [personal["school"], personal["email"], personal["phone"], personal["location"], personal["website"]]))
    if contacts:
        story.append(Paragraph(pdf_text(contacts), meta_style))

    for section in data["sections"]:
        story.append(Spacer(1, 2))
        story.append(Paragraph(pdf_text(section["title"]), heading_style))
        if section["type"] == "text":
            story.append(Paragraph(pdf_text(section["items"][0]["text"]), body_style))
            continue
        for item in section["items"]:
            line_parts = [f"<b>{pdf_inline(item['title'])}</b>" if item["title"] else ""]
            if item["subtitle"]:
                line_parts.append(pdf_inline(item["subtitle"]))
            if item["meta"]:
                line_parts.append(f"<font color='{accent}'>{pdf_inline(item['meta'])}</font>")
            line_text = " &nbsp;&nbsp; ".join(part for part in line_parts if part)
            if line_text:
                story.append(Paragraph(line_text, body_style))
            if item["detail"]:
                story.append(Paragraph(pdf_text(item["detail"]), body_style))
            for bullet in item["bullets"]:
                story.append(Paragraph(pdf_text(bullet), bullet_style, bulletText="•"))

    doc.build(story)


def load_pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if FONT_PATH:
        candidates.append(FONT_PATH)
    candidates.extend(
        [
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",
        ]
    )
    for candidate in candidates:
        if os.path.exists(candidate):
            try:
                return ImageFont.truetype(candidate, size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def template_background(template_id: str) -> tuple[str, str]:
    mapping = {
        "harbor": ("#f5f7fb", "#ffffff"),
        "editorial": ("#f7f2ec", "#fffdfa"),
        "jade": ("#f2faf7", "#fbfffd"),
    }
    return mapping.get(template_id, ("#f5f7fb", "#ffffff"))


def draw_wrapped_text(draw: ImageDraw.ImageDraw, text: str, font, fill: str, x: int, y: int, width: int, line_gap: int) -> int:
    words = list(text)
    line = ""
    for token in words:
        trial = line + token
        box = draw.textbbox((0, 0), trial, font=font)
        if box[2] - box[0] > width and line:
            draw.text((x, y), line, font=font, fill=fill)
            y += (box[3] - box[1]) + line_gap
            line = token
        else:
            line = trial
    if line:
        box = draw.textbbox((0, 0), line, font=font)
        draw.text((x, y), line, font=font, fill=fill)
        y += (box[3] - box[1]) + line_gap
    return y


def build_png(data: dict, output_path: Path) -> None:
    width, height = 1654, 2339
    outer_bg, panel_bg = template_background(data["templateId"])
    accent = data["theme"]["accent"]
    scale = data["theme"]["fontScale"]
    image = Image.new("RGB", (width, height), outer_bg)
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((90, 70, width - 90, height - 70), radius=42, fill=panel_bg)
    draw.rectangle((90, 70, 130, height - 70), fill=accent)

    name_font = load_pil_font(int(52 * scale))
    title_font = load_pil_font(int(24 * scale))
    section_font = load_pil_font(int(25 * scale))
    body_font = load_pil_font(int(18 * scale))
    small_font = load_pil_font(int(16 * scale))

    personal = data["personal"]
    y = 120
    draw.text((160, y), personal["name"], font=name_font, fill="#111111")
    y += 80
    if personal["tagline"]:
        draw.text((160, y), personal["tagline"], font=title_font, fill=accent)
        y += 42
    contacts = "  |  ".join(filter(None, [personal["school"], personal["email"], personal["phone"], personal["location"], personal["website"]]))
    if contacts:
        y = draw_wrapped_text(draw, contacts, small_font, "#666666", 160, y, width - 320, 8)
        y += 18

    for section in data["sections"]:
        draw.text((160, y), section["title"].upper(), font=section_font, fill=accent)
        y += 40
        draw.line((160, y, width - 150, y), fill=accent, width=2)
        y += 18
        if section["type"] == "text":
            y = draw_wrapped_text(draw, section["items"][0]["text"], body_font, "#222222", 160, y, width - 310, 10)
            y += 12
            continue
        for item in section["items"]:
            if item["title"]:
                title_line = item["title"]
                if item["subtitle"]:
                    title_line += f"  |  {item['subtitle']}"
                if item["meta"]:
                    title_line += f"  |  {item['meta']}"
                y = draw_wrapped_text(draw, title_line, body_font, "#111111", 160, y, width - 310, 8)
            if item["detail"]:
                y = draw_wrapped_text(draw, item["detail"], small_font, "#444444", 160, y, width - 310, 6)
            for bullet in item["bullets"]:
                draw.text((180, y), "•", font=body_font, fill=accent)
                y = draw_wrapped_text(draw, bullet, small_font, "#333333", 210, y, width - 370, 6)
            y += 10
        y += 10
        if y > height - 180:
            break

    image.save(output_path, format="PNG")


def build_with_chrome(data: dict, output_path: Path, export_format: str) -> None:
    node_executable = find_node_executable()
    if not node_executable:
        raise RuntimeError("未找到 Node.js，无法启动 Chrome 导出器")
    if not CHROME_RENDERER.exists():
        raise RuntimeError("Chrome 导出器文件缺失")

    with tempfile.TemporaryDirectory(prefix="resume-export-") as temp_dir:
        payload_path = Path(temp_dir) / "resume.json"
        payload_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        result = subprocess.run(
            [
                str(node_executable),
                str(CHROME_RENDERER),
                export_format,
                str(payload_path),
                str(output_path),
            ],
            cwd=str(ROOT),
            capture_output=True,
            text=True,
            timeout=45,
            check=False,
        )

    if result.returncode != 0:
        detail = (result.stderr or result.stdout or "Chrome 渲染器启动失败").strip()
        raise RuntimeError(detail[-800:])
    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("Chrome 渲染器没有生成有效文件")


def build_visual_export(data: dict, output_path: Path, export_format: str) -> tuple[str, str | None]:
    try:
        build_with_chrome(data, output_path, export_format)
        return "Chrome", None
    except Exception as chrome_error:
        if output_path.exists():
            output_path.unlink()
        if export_format == "pdf":
            build_pdf(data, output_path)
        else:
            build_png(data, output_path)
        return "Python fallback", str(chrome_error)


class ResumeHandler(BaseHTTPRequestHandler):
    def _send_json(self, payload: dict, status: int = 200) -> None:
        body = json.dumps(payload).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _send_file(self, path: Path) -> None:
        if not path.exists() or not path.is_file():
            self.send_error(HTTPStatus.NOT_FOUND, "File not found")
            return
        mime = {
            ".html": "text/html; charset=utf-8",
            ".css": "text/css; charset=utf-8",
            ".js": "application/javascript; charset=utf-8",
            ".json": "application/json; charset=utf-8",
            ".png": "image/png",
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".svg": "image/svg+xml",
        }.get(path.suffix.lower(), "application/octet-stream")
        raw = path.read_bytes()
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", mime)
        if path.parent == EXPORT_DIR:
            encoded_name = quote(path.name)
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{encoded_name}")
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def do_GET(self):
        parsed = urlparse(self.path)
        if parsed.path == "/":
            return self._send_file(STATIC_DIR / "index.html")
        if parsed.path == "/api/health":
            return self._send_json(
                {
                    "ok": True,
                    "chromeRenderer": bool(find_node_executable() and CHROME_RENDERER.exists()),
                    "exportsDir": str(EXPORT_DIR),
                }
            )
        if parsed.path.startswith("/exports/"):
            export_name = Path(unquote(parsed.path)).name
            return self._send_file(EXPORT_DIR / export_name)
        safe_path = (STATIC_DIR / unquote(parsed.path).lstrip("/")).resolve()
        if STATIC_DIR.resolve() not in safe_path.parents and safe_path != STATIC_DIR.resolve():
            self.send_error(HTTPStatus.FORBIDDEN, "Forbidden")
            return
        return self._send_file(safe_path)

    def do_POST(self):
        parsed = urlparse(self.path)
        if parsed.path not in {"/api/export/pdf", "/api/export/docx", "/api/export/png"}:
            self.send_error(HTTPStatus.NOT_FOUND, "Unknown endpoint")
            return

        length = int(self.headers.get("Content-Length", "0"))
        if not length:
            return self._send_json({"error": "Missing body"}, 400)
        try:
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
        except json.JSONDecodeError:
            return self._send_json({"error": "Invalid JSON"}, 400)

        data = normalize_payload(payload)
        ensure_export_dir()
        stem = build_safe_stem(data["personal"]["name"])
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        engine = ""
        warning = None
        try:
            if parsed.path.endswith("/pdf"):
                target = EXPORT_DIR / f"{stem}-{stamp}.pdf"
                engine, warning = build_visual_export(data, target, "pdf")
            elif parsed.path.endswith("/docx"):
                target = EXPORT_DIR / f"{stem}-{stamp}.docx"
                build_docx(data, target)
                engine = "python-docx"
            else:
                target = EXPORT_DIR / f"{stem}-{stamp}.png"
                engine, warning = build_visual_export(data, target, "png")
        except Exception as exc:
            return self._send_json({"ok": False, "error": f"导出失败：{exc}"}, 500)

        return self._send_json(
            {
                "ok": True,
                "filename": target.name,
                "path": f"/exports/{target.name}",
                "engine": engine,
                "warning": warning,
            }
        )


def main() -> None:
    ensure_export_dir()
    server = ThreadingHTTPServer((HOST, PORT), ResumeHandler)
    print(f"Resume generator is running at http://{HOST}:{PORT}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nServer stopped.")


if __name__ == "__main__":
    main()
